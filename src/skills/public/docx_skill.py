"""DOCX generation skill -- converts various input formats into Word documents."""

from __future__ import annotations

import json
import re
import traceback

from src.generators.base import (
    ContentBlock,
    ContentType,
    GenerationMetadata,
    SectionData,
    StructuredContent,
)
from src.skills.base import BaseSkill
from src.skills.models import QACheck, QAReport, SkillMetadata, SkillResult
from src.utils.logging import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Content format detection
# ---------------------------------------------------------------------------

def _detect_format(content: str) -> str:
    """Heuristically detect the format of *content*."""
    stripped = content.strip()

    if stripped.startswith("{") or stripped.startswith("["):
        try:
            json.loads(stripped)
            return "json"
        except (json.JSONDecodeError, ValueError):
            pass

    first_line = stripped.split("\n", 1)[0]
    if "," in first_line and "<" not in first_line and first_line.count(",") >= 2:
        return "csv"

    if "<html" in stripped.lower() or stripped.startswith("<!") or "</" in stripped[:200]:
        return "html"

    if any(stripped.startswith(c) for c in ("# ", "## ", "### ")):
        return "markdown"
    if "\n# " in stripped or "\n## " in stripped or "\n- " in stripped or "\n* " in stripped:
        return "markdown"

    return "text"


# ---------------------------------------------------------------------------
# Lightweight parsers -> StructuredContent (sections-based)
# ---------------------------------------------------------------------------

def _parse_markdown_to_sections(content: str, title: str) -> StructuredContent:
    """Turn Markdown text into sections split on ``##`` headings."""
    sections: list[SectionData] = []
    current_title = ""
    current_blocks: list[ContentBlock] = []
    doc_title = title

    for line in content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = stripped.lstrip("# ").strip()
            if not doc_title:
                doc_title = heading
            continue

        if stripped.startswith("## "):
            if current_title or current_blocks:
                sections.append(SectionData(title=current_title, blocks=current_blocks))
            current_title = stripped.lstrip("# ").strip()
            current_blocks = []
            continue

        if stripped.startswith("### "):
            level = len(stripped.split()[0])
            current_blocks.append(
                ContentBlock(type=ContentType.HEADING, content=stripped.lstrip("# ").strip(), level=level)
            )
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:].strip()
            if current_blocks and current_blocks[-1].type == ContentType.BULLET_LIST:
                current_blocks[-1].items.append(item_text)
            else:
                current_blocks.append(ContentBlock(type=ContentType.BULLET_LIST, items=[item_text]))
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            item_text = re.sub(r"^\d+\.\s*", "", stripped)
            if current_blocks and current_blocks[-1].type == ContentType.NUMBERED_LIST:
                current_blocks[-1].items.append(item_text)
            else:
                current_blocks.append(ContentBlock(type=ContentType.NUMBERED_LIST, items=[item_text]))
            continue

        if stripped:
            current_blocks.append(ContentBlock(type=ContentType.TEXT, content=stripped))

    if current_title or current_blocks:
        sections.append(SectionData(title=current_title, blocks=current_blocks))

    return StructuredContent(title=doc_title or "Untitled", sections=sections)


def _parse_text_to_sections(content: str, title: str) -> StructuredContent:
    """Split plain text into one section per paragraph."""
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    blocks = [ContentBlock(type=ContentType.TEXT, content=p) for p in paragraphs]
    section = SectionData(title="", blocks=blocks)
    return StructuredContent(title=title or "Untitled", sections=[section])


def _parse_json_to_sections(content: str, title: str) -> StructuredContent:
    data = json.loads(content)
    if isinstance(data, dict) and ("sections" in data or "slides" in data):
        sc = StructuredContent.model_validate(data)
        if title:
            sc.title = title
        return sc
    if isinstance(data, list):
        sections = []
        for item in data:
            sec_title = item.get("title", "")
            body = item.get("body", item.get("content", ""))
            blocks = [ContentBlock(type=ContentType.TEXT, content=body)] if body else []
            sections.append(SectionData(title=sec_title, blocks=blocks))
        return StructuredContent(title=title or "Untitled", sections=sections)
    return _parse_text_to_sections(json.dumps(data, indent=2), title)


def _parse_csv_to_sections(content: str, title: str) -> StructuredContent:
    rows: list[list[str]] = []
    for line in content.strip().split("\n"):
        if line.strip():
            rows.append([cell.strip() for cell in line.split(",")])
    table_block = ContentBlock(type=ContentType.TABLE, rows=rows)
    section = SectionData(title=title or "Data", blocks=[table_block])
    return StructuredContent(title=title or "Data", sections=[section])


def _parse_html_to_sections(content: str, title: str) -> StructuredContent:
    text = re.sub(r"<[^>]+>", " ", content)
    text = re.sub(r"\s+", " ", text).strip()
    return _parse_text_to_sections(text, title)


# ---------------------------------------------------------------------------
# The skill
# ---------------------------------------------------------------------------

class DocxSkill(BaseSkill):
    """Skill that generates DOCX (Word) documents."""

    @property
    def metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="docx",
            description="Generate Word documents from various input formats",
            input_formats=["html", "markdown", "json", "text", "csv"],
            output_format="docx",
            version="1.0.0",
            tags=["document", "office", "docx"],
        )

    async def validate_params(self, params: dict) -> bool:
        if "content" not in params:
            return False
        if not isinstance(params["content"], str):
            return False
        if not params["content"].strip():
            return False
        return True

    async def execute(self, params: dict, context: dict) -> SkillResult:
        try:
            content_str: str = params["content"]
            content_format: str = params.get("content_format", "") or _detect_format(content_str)
            title: str = params.get("title", "")
            template: str = params.get("template", "default")
            style: dict = params.get("style", {})

            parser_map = {
                "markdown": _parse_markdown_to_sections,
                "text": _parse_text_to_sections,
                "json": _parse_json_to_sections,
                "csv": _parse_csv_to_sections,
                "html": _parse_html_to_sections,
            }
            parser = parser_map.get(content_format, _parse_text_to_sections)
            structured = parser(content_str, title)

            gen_meta = GenerationMetadata(template=template, style=style)

            try:
                from src.generators.docx_generator import DocxGenerator
                generator = DocxGenerator()
            except ImportError:
                logger.warning("docx_generator_not_found", detail="Using fallback generator")
                generator = _FallbackDocxGenerator()

            output_bytes = await generator.generate(structured, gen_meta)

            return SkillResult(
                success=True,
                output_format="docx",
                content_bytes=output_bytes,
                metadata={
                    "title": structured.title,
                    "section_count": len(structured.sections),
                    "content_format": content_format,
                    "template": template,
                },
            )

        except Exception as exc:
            logger.exception("docx_skill_error")
            return SkillResult(
                success=False,
                output_format="docx",
                error=f"{type(exc).__name__}: {exc}",
                metadata={"traceback": traceback.format_exc()},
            )

    async def qa_check(self, result: SkillResult) -> QAReport:
        report = await super().qa_check(result)
        if result.content_bytes:
            is_zip = result.content_bytes[:4] == b"PK\x03\x04"
            report.checks.append(
                QACheck(
                    name="valid_docx_header",
                    passed=is_zip,
                    detail="File starts with PK zip header" if is_zip else "Invalid DOCX header",
                )
            )
            report.passed = report.passed and is_zip
        return report


# ---------------------------------------------------------------------------
# Fallback generator
# ---------------------------------------------------------------------------

class _FallbackDocxGenerator:
    """Minimal DOCX generator using python-docx directly."""

    async def generate(self, content: StructuredContent, metadata: GenerationMetadata) -> bytes:
        from io import BytesIO

        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Title
        if content.title:
            doc.add_heading(content.title, level=0)

        if content.subtitle:
            doc.add_paragraph(content.subtitle)

        for section in content.sections:
            if section.title:
                doc.add_heading(section.title, level=1)

            for block in section.blocks:
                if block.type == ContentType.HEADING:
                    doc.add_heading(block.content, level=min(block.level, 4))

                elif block.type == ContentType.TEXT:
                    doc.add_paragraph(block.content)

                elif block.type == ContentType.BULLET_LIST:
                    for item in block.items:
                        doc.add_paragraph(item, style="List Bullet")

                elif block.type == ContentType.NUMBERED_LIST:
                    for item in block.items:
                        doc.add_paragraph(item, style="List Number")

                elif block.type == ContentType.TABLE and block.rows:
                    rows = block.rows
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        for j, cell_text in enumerate(row):
                            if j < len(table.columns):
                                table.cell(i, j).text = cell_text

                elif block.type == ContentType.CODE:
                    p = doc.add_paragraph()
                    run = p.add_run(block.content)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
