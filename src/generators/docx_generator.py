"""Word document generator using python-docx."""

import io

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from .base import (
    BaseGenerator,
    ContentBlock,
    ContentType,
    GenerationMetadata,
    SectionData,
    StructuredContent,
)


DEFAULT_COLORS = {
    "primary": "2B579A",
    "secondary": "217346",
    "accent": "B7472A",
    "text": "333333",
    "table_header": "2B579A",
    "table_alt_row": "F2F2F2",
}

DEFAULT_FONTS = {
    "title": "Calibri",
    "body": "Calibri",
}


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex string to RGBColor."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _resolve_colors(metadata: GenerationMetadata) -> dict[str, str]:
    colors = dict(DEFAULT_COLORS)
    colors.update(metadata.style.get("color_scheme", {}))
    return colors


def _resolve_fonts(metadata: GenerationMetadata) -> dict[str, str]:
    fonts = dict(DEFAULT_FONTS)
    fonts.update(metadata.style.get("fonts", {}))
    return fonts


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


class DOCXGenerator(BaseGenerator):
    """Generates Word (.docx) documents from structured content."""

    async def generate(
        self, content: StructuredContent, metadata: GenerationMetadata
    ) -> bytes:
        """Generate a DOCX file and return it as bytes."""
        doc = Document()
        colors = _resolve_colors(metadata)
        fonts = _resolve_fonts(metadata)

        self._configure_default_style(doc, colors, fonts)

        # Title page
        if content.title:
            self._add_title_page(doc, content, colors, fonts)

        # Render each section
        for idx, section in enumerate(content.sections):
            if idx > 0 or content.title:
                # Page break before each new section (except first if no title page)
                doc.add_page_break()
            self._render_section(doc, section, colors, fonts)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------ #
    #  Document setup
    # ------------------------------------------------------------------ #

    def _configure_default_style(self, doc: Document, colors: dict, fonts: dict):
        """Set up the default paragraph and heading styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = fonts["body"]
        font.size = Pt(11)
        font.color.rgb = _hex_to_rgb(colors["text"])

        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)

        # Configure heading styles
        for level in range(1, 5):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_font = h_style.font
                h_font.name = fonts["title"]
                h_font.bold = True
                h_font.color.rgb = _hex_to_rgb(colors["primary"])
                sizes = {1: 24, 2: 20, 3: 16, 4: 14}
                h_font.size = Pt(sizes.get(level, 14))

    # ------------------------------------------------------------------ #
    #  Title page
    # ------------------------------------------------------------------ #

    def _add_title_page(
        self, doc: Document, content: StructuredContent, colors: dict, fonts: dict
    ):
        """Create a title page with title, subtitle, and author."""
        # Add some vertical spacing
        for _ in range(6):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(content.title)
        title_run.font.name = fonts["title"]
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = _hex_to_rgb(colors["primary"])

        # Subtitle
        if content.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(12)
            sub_run = sub_para.add_run(content.subtitle)
            sub_run.font.name = fonts["body"]
            sub_run.font.size = Pt(18)
            sub_run.font.color.rgb = _hex_to_rgb(colors.get("subtitle", "666666"))

        # Author
        if content.author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.paragraph_format.space_before = Pt(24)
            author_run = author_para.add_run(content.author)
            author_run.font.name = fonts["body"]
            author_run.font.size = Pt(14)
            author_run.font.color.rgb = _hex_to_rgb(colors["text"])
            author_run.font.italic = True

    # ------------------------------------------------------------------ #
    #  Section rendering
    # ------------------------------------------------------------------ #

    def _render_section(
        self, doc: Document, section: SectionData, colors: dict, fonts: dict
    ):
        """Render a complete section with its heading and content blocks."""
        if section.title:
            doc.add_heading(section.title, level=1)

        for block in section.blocks:
            self._render_block(doc, block, colors, fonts)

    def _render_block(
        self, doc: Document, block: ContentBlock, colors: dict, fonts: dict
    ):
        """Render a single content block."""
        if block.type == ContentType.HEADING:
            self._render_heading(doc, block, colors, fonts)
        elif block.type == ContentType.TEXT:
            self._render_text(doc, block, colors, fonts)
        elif block.type == ContentType.BULLET_LIST:
            self._render_bullet_list(doc, block, colors, fonts)
        elif block.type == ContentType.NUMBERED_LIST:
            self._render_numbered_list(doc, block, colors, fonts)
        elif block.type == ContentType.TABLE:
            self._render_table(doc, block, colors, fonts)
        elif block.type == ContentType.CODE:
            self._render_code(doc, block, colors, fonts)
        elif block.type == ContentType.IMAGE:
            self._render_image_placeholder(doc, block, colors, fonts)
        else:
            # Fallback: render as plain text
            if block.content:
                doc.add_paragraph(block.content)

    def _render_heading(self, doc, block, colors, fonts):
        """Render a heading block with the appropriate level (1-4)."""
        level = max(1, min(block.level, 4))
        doc.add_heading(block.content, level=level)

    def _render_text(self, doc, block, colors, fonts):
        """Render a plain text paragraph."""
        if not block.content:
            return
        para = doc.add_paragraph()
        run = para.add_run(block.content)
        run.font.name = fonts["body"]
        run.font.size = Pt(11)
        run.font.color.rgb = _hex_to_rgb(colors["text"])

    def _render_bullet_list(self, doc, block, colors, fonts):
        """Render a bullet (unordered) list."""
        for item in block.items:
            para = doc.add_paragraph(style="List Bullet")
            # Clear existing text and apply font styling
            para.clear()
            run = para.add_run(item)
            run.font.name = fonts["body"]
            run.font.size = Pt(11)
            run.font.color.rgb = _hex_to_rgb(colors["text"])

    def _render_numbered_list(self, doc, block, colors, fonts):
        """Render a numbered (ordered) list."""
        for item in block.items:
            para = doc.add_paragraph(style="List Number")
            para.clear()
            run = para.add_run(item)
            run.font.name = fonts["body"]
            run.font.size = Pt(11)
            run.font.color.rgb = _hex_to_rgb(colors["text"])

    def _render_table(self, doc, block, colors, fonts):
        """Render a table with header formatting and alternating row colors."""
        rows_data = block.rows
        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(rows_data):
            row = table.rows[row_idx]
            for col_idx in range(num_cols):
                cell = row.cells[col_idx]
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell.text = ""

                para = cell.paragraphs[0]
                run = para.add_run(str(cell_text))

                if row_idx == 0:
                    # Header row
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = fonts["body"]
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_shading(cell, colors["table_header"])
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run.font.size = Pt(10)
                    run.font.name = fonts["body"]
                    run.font.color.rgb = _hex_to_rgb(colors["text"])
                    # Alternating row colors
                    if row_idx % 2 == 0:
                        _set_cell_shading(cell, colors["table_alt_row"])

        # Add spacing after table
        doc.add_paragraph()

    def _render_code(self, doc, block, colors, fonts):
        """Render a code block with monospace font and gray background."""
        if not block.content:
            return

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # Apply a light gray shading to the paragraph via XML
        pPr = para._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
        )
        pPr.append(shading)

        run = para.add_run(block.content)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb("333333")

    def _render_image_placeholder(self, doc, block, colors, fonts):
        """Render a placeholder for an image (since we can't embed without the actual file)."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[Image: {block.content or block.metadata.get('url', 'image')}]")
        run.font.italic = True
        run.font.color.rgb = _hex_to_rgb("999999")
        run.font.size = Pt(10)
